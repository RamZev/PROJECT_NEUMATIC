# neumatic\apps\informes\views\cliente_list_views.py
from django.urls import reverse_lazy
from ..views.list_views_generics import *
from apps.maestros.models.cliente_models import Cliente
from apps.maestros.models.base_models import *
from django.utils import timezone

from django.http import HttpResponse
from django.views import View
from reportlab.lib.pagesizes import letter
from io import BytesIO
from zipfile import ZipFile
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle,Paragraph, Spacer
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from docx import Document
from openpyxl import Workbook
import csv
import io
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle


class ConfigViews:
    # Modelo
    model = Cliente
    
    # Aplicación asociada al modelo
    app_label = "informes"
    
    # Nombre del modelo en minúsculas
    model_string = model.__name__.lower()
    
    # Vistas del CRUD del modelo
    list_view_name = f"{model_string}_list"
    
    # Plantilla de la lista del CRUD
    template_list = f'{app_label}/maestro_informe_list.html'
    
    # Contexto de los datos de la lista
    context_object_name = 'objetos'
    
    # Vista del home del proyecto
    home_view_name = "home"
    
    # Nombre de la url 
    success_url = reverse_lazy(list_view_name)


class DataViewList:
    search_fields = ['nombre_cliente', 'cuit']
    
    ordering = ['nombre_cliente']
    
    paginate_by = 8
	  
    table_headers = {
        'estatus_cliente': (2, 'Estatus'),
        'nombre_cliente': (4, 'Nombre Cliente'),
        'tipo_persona': (2, 'Tipo'),
        'cuit': (2, 'CUIT'),
    }
    
    table_data = [
        {'field_name': 'estatus_cliente', 'date_format': None},
        {'field_name': 'nombre_cliente', 'date_format': None},
        {'field_name': 'tipo_persona', 'date_format': None},
        {'field_name': 'cuit', 'date_format': None},
    ]


# ClienteListView - Inicio
class ClienteInformeListView(InformeListView):
    model = ConfigViews.model
    template_name = ConfigViews.template_list
    context_object_name = ConfigViews.context_object_name
    
    search_fields = DataViewList.search_fields
    ordering = DataViewList.ordering
    
    extra_context = {
        "master_title": f'Informes - {ConfigViews.model._meta.verbose_name_plural}',
        "home_view_name": ConfigViews.home_view_name,
        "list_view_name": ConfigViews.list_view_name,
        "table_headers": DataViewList.table_headers,
        "table_data": DataViewList.table_data,
        "buscador_template": "informes/buscador_cliente.html",  # Se agrega el nombre de la plantilla buscador
    }

    def get_queryset(self):
        queryset = super().get_queryset()
        
        # Filtrar según los campos personalizados
        estatus_cliente = self.request.GET.get('estatus_cliente')
        nombre_cliente = self.request.GET.get('nombre_cliente')
        tipo_persona = self.request.GET.get('tipo_persona')
        
        if nombre_cliente:
            queryset = queryset.filter(nombre_cliente=nombre_cliente)

        if estatus_cliente:
            queryset = queryset.filter(estatus_cliente=estatus_cliente)
        
        if tipo_persona:
            queryset = queryset.filter(tipo_persona=tipo_persona)
        
        return queryset


class ClienteInformesView(View):
    def get(self, request, *args, **kwargs):
        # Obtener el queryset (puedes modificar este método según tu lógica)
        cliente_list_view = ClienteInformeListView()
        cliente_list_view.request = request
        queryset = cliente_list_view.get_queryset()
        
        
        # Generar y retornar el archivo ZIP
        return self.generar_archivos_zip(queryset)

    def generar_archivos_zip(self, queryset):
        buffer = BytesIO()
        with ZipFile(buffer, 'w') as zip_file:
            # Generar y agregar cada archivo al ZIP
            self.agregar_pdf(queryset, zip_file)
            self.agregar_csv(queryset, zip_file)
            self.agregar_word(queryset, zip_file)
            self.agregar_excel(queryset, zip_file)

        # Preparar respuesta para descargar el archivo ZIP
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/zip')
        response['Content-Disposition'] = 'attachment; filename="informe_clientes.zip"'
        return response


    def agregar_pdf(self, queryset, zip_file):
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)

        # Crear un canvas para agregar el número de página y la línea
        def agregar_numero_pagina(canvas, doc):
            # Dibujar una línea horizontal antes del número de página
            canvas.setStrokeColor(colors.black)  # Color de la línea
            canvas.setLineWidth(0.5)  # Grosor de la línea
            canvas.line(40, 30, 570, 30)  # Coordenadas para la línea horizontal (ajustar si es necesario)

            # Dibujar el número de página
            canvas.setFont("Helvetica", 8)
            page_num = canvas.getPageNumber()
            canvas.drawRightString(580, 10, f"Página {page_num}")

        doc.build([], onFirstPage=agregar_numero_pagina, onLaterPages=agregar_numero_pagina)

        # Crear título, subtítulo, tabla y elementos, como lo hacías antes
        styles = getSampleStyleSheet()
        title_style = styles['Title']
        centered_subtitle_style = ParagraphStyle(
            name='CenteredSubtitle',
            parent=styles['Normal'],
            alignment=1  # 1 indica centrado
        )
        
        # Crear título y subtítulo con líneas debajo
        title = Paragraph("Título de Prueba", title_style)
        subtitle = Paragraph("Listado Prueba", centered_subtitle_style)
        line = Spacer(1, 12)
        line_below = Table([['']], colWidths=[450])  # Línea decorativa
        line_below.setStyle(TableStyle([('LINEBELOW', (0, 0), (-1, -1), 1, colors.black)]))
        
        # Crear datos para la tabla, alineando texto y números
        data = [['Nombre Cliente', 'CUIT', 'Estatus', 'Tipo Persona']] + [
            [
                obj.nombre_cliente,  # Texto
                Paragraph(f"<para align='right'>{obj.cuit}</para>", styles['Normal']),  # Número
                obj.estatus_cliente,  # Texto
                obj.tipo_persona  # Texto
            ]
            for obj in queryset
        ]
        
        table = Table(data, colWidths=[150, 100, 100, 100])  # Ajustar anchos de columnas
        table.setStyle(TableStyle([ 
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 1), (0, -1), 'LEFT'),  # Alinear a la izquierda columna de texto
            ('ALIGN', (1, 1), (1, -1), 'RIGHT'),  # Alinear a la derecha columna de números
            ('ALIGN', (2, 1), (-1, -1), 'LEFT'),  # Alinear las demás columnas a la izquierda
        ]))
        
        # Construir el documento con título, línea decorativa, subtítulo y tabla
        elements = [
            title, line_below, Spacer(1, 12),  # Título y línea
            subtitle, line_below, Spacer(1, 12),  # Subtítulo y línea
            table
        ]
        
        doc.build(elements, onFirstPage=agregar_numero_pagina, onLaterPages=agregar_numero_pagina)
        
        # Guardar el PDF en el archivo zip
        zip_file.writestr('informe_clientes.pdf', buffer.getvalue())
        buffer.close()

    def agregar_csv(self, queryset, zip_file):
        # Crear un buffer en memoria
        buffer = BytesIO()
        text_buffer = io.TextIOWrapper(buffer, encoding='utf-8', newline='')

        # Escribir el contenido CSV
        writer = csv.writer(text_buffer)
        writer.writerow(['Nombre Cliente', 'CUIT', 'Estatus', 'Tipo Persona'])
        for obj in queryset:
            writer.writerow([obj.nombre_cliente, obj.cuit, obj.estatus_cliente, obj.tipo_persona])

        # Asegurar que el contenido esté escrito y el cursor en el lugar correcto
        text_buffer.flush()  # Asegurar que se escriba el contenido al buffer
        buffer.seek(0)

        # Guardar en el archivo ZIP
        zip_file.writestr('informe_clientes.csv', buffer.read())

        # Cerrar ambos buffers
        text_buffer.close()
        buffer.close()


    def agregar_word(self, queryset, zip_file):
        buffer = BytesIO()
        doc = Document()
        doc.add_heading('Informe de Clientes', level=1)
        
        # Crear una tabla con encabezados
        table = doc.add_table(rows=1, cols=4)
        headers = ['Nombre Cliente', 'CUIT', 'Estatus', 'Tipo Persona']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        
        # Añadir las filas con datos del queryset
        for obj in queryset:
            row_cells = table.add_row().cells
            row_cells[0].text = str(obj.nombre_cliente)  # Convertir a cadena
            row_cells[1].text = str(obj.cuit)           # Convertir a cadena
            row_cells[2].text = str(obj.estatus_cliente) # Convertir a cadena
            row_cells[3].text = str(obj.tipo_persona)   # Convertir a cadena

        # Guardar el documento en el buffer
        doc.save(buffer)
        buffer.seek(0)
        zip_file.writestr('informe_clientes.docx', buffer.getvalue())
        buffer.close()


    def agregar_excel(self, queryset, zip_file):
        buffer = BytesIO()
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = 'Informe de Clientes'
        headers = ['Nombre Cliente', 'CUIT', 'Estatus', 'Tipo Persona']
        sheet.append(headers)
        for obj in queryset:
            sheet.append([obj.nombre_cliente, obj.cuit, obj.estatus_cliente, obj.tipo_persona])
        workbook.save(buffer)
        buffer.seek(0)
        zip_file.writestr('informe_clientes.xlsx', buffer.getvalue())
        buffer.close()


class ClienteInformePDFView(View):
    def get(self, request, *args, **kwargs):
        # Obtener el queryset (el listado de clientes)
        cliente_list_view = ClienteInformeListView()  # Suponiendo que esta vista genera el queryset
        cliente_list_view.request = request
        queryset = cliente_list_view.get_queryset()

        # Crear PDF en memoria
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)

        # Función para agregar número de página
        def agregar_numero_pagina(canvas, doc):
            # Dibujar una línea horizontal antes del número de página
            canvas.setStrokeColor(colors.black)  # Color de la línea
            canvas.setLineWidth(0.5)  # Grosor de la línea
            canvas.line(40, 30, 570, 30)  # Coordenadas para la línea horizontal (ajustar si es necesario)

            # Dibujar el número de página
            canvas.setFont("Helvetica", 8)
            page_num = canvas.getPageNumber()
            canvas.drawRightString(580, 10, f"Página {page_num}")

        # Estilos para el documento
        styles = getSampleStyleSheet()
        title_style = styles['Title']
        centered_subtitle_style = ParagraphStyle(
            name='CenteredSubtitle',
            parent=styles['Normal'],
            alignment=1  # Centrado
        )

        # Crear título y subtítulo con líneas debajo
        title = Paragraph("Título de Prueba", title_style)
        subtitle = Paragraph("Listado Prueba", centered_subtitle_style)
        line = Spacer(1, 12)
        line_below = Table([['']], colWidths=[450])  # Línea decorativa
        line_below.setStyle(TableStyle([('LINEBELOW', (0, 0), (-1, -1), 1, colors.black)]))

        # Crear datos para la tabla, alineando texto y números
        data = [['Nombre Cliente', 'CUIT', 'Estatus', 'Tipo Persona']] + [
            [
                obj.nombre_cliente,  # Texto
                Paragraph(f"<para align='right'>{obj.cuit}</para>", styles['Normal']),  # Número
                obj.estatus_cliente,  # Texto
                obj.tipo_persona  # Texto
            ]
            for obj in queryset
        ]

        # Crear la tabla
        table = Table(data, colWidths=[150, 100, 100, 100])  # Ajustar anchos de columnas
        table.setStyle(TableStyle([ 
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 1), (0, -1), 'LEFT'),  # Alinear a la izquierda columna de texto
            ('ALIGN', (1, 1), (1, -1), 'RIGHT'),  # Alinear a la derecha columna de números
            ('ALIGN', (2, 1), (-1, -1), 'LEFT'),  # Alinear las demás columnas a la izquierda
        ]))

        # Construir el documento con título, línea decorativa, subtítulo y tabla
        elements = [
            title, line_below, Spacer(1, 12),  # Título y línea
            subtitle, line_below, Spacer(1, 12),  # Subtítulo y línea
            table
        ]

        # Construir el PDF
        doc.build(elements, onFirstPage=agregar_numero_pagina, onLaterPages=agregar_numero_pagina)

        # Preparar la respuesta HTTP para mostrar el PDF
        buffer.seek(0)
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'inline; filename="informe_clientes.pdf"'
        return response