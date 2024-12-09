# neumatic\apps\informes\views\cliente_list_views.py
from django.urls import reverse_lazy
from ..views.list_views_generics import *
from apps.maestros.models.cliente_models import Cliente
from ..forms.buscador_cliente_forms import BuscadorClienteForm

# from apps.maestros.models.base_models import *
# from django.utils import timezone

from django.http import HttpResponse, JsonResponse
from django.views import View
from reportlab.lib.pagesizes import letter
from io import BytesIO, TextIOWrapper
from zipfile import ZipFile
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from docx import Document
from openpyxl import Workbook
import csv
import io
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from django.core.mail import EmailMessage

from django.db.models.functions import Lower
from django.db.models import Q

class ConfigViews:
	# Modelo
	model = Cliente
	
	# Formulario asociado al modelo
	form_class = BuscadorClienteForm
	
	# Aplicación asociada al modelo
	app_label = "informes"
	
	# Nombre del modelo en minúsculas
	model_string = model.__name__.lower()
	
	# Vistas del CRUD del modelo
	list_view_name = f"{model_string}_list"
	
	# Plantilla de la lista del CRUD
	template_list = f'{app_label}/maestro_informe_list.html'
	
	# Contexto de los datos de la lista
	context_object_name = 'objetos'
	
	# Vista del home del proyecto
	home_view_name = "home"
	
	# Nombre de la url 
	success_url = reverse_lazy(list_view_name)
	
	# Archivo JavaScript específico.
	js_file = 'js/filtro_cliente.js'


class DataViewList:
	search_fields = ['nombre_cliente', 'cuit']
	
	ordering = ['nombre_cliente']
	
	paginate_by = 8
	  
	table_headers = {
		'id_cliente': (1, 'Código'),
		'nombre_cliente': (3, 'Nombre Cliente'),
		'domicilio_cliente': (3, 'Domicilio'),
		'id_localidad': (1, 'Localidad'),
		'id_localidad.codigo_postal': (1, 'C.P.'),
		'id_tipo_iva.codigo_iva': (1, 'IVA'),
		'cuit': (1, 'CUIT'),
		'telefono_cliente': (1, 'Teléfono'),
	}
	
	table_data = [
		{'field_name': 'id_cliente', 'date_format': None},
		{'field_name': 'nombre_cliente', 'date_format': None},
		{'field_name': 'domicilio_cliente', 'date_format': None},
		{'field_name': 'id_localidad', 'date_format': None},
		{'field_name': 'id_localidad.codigo_postal', 'date_format': None},
		{'field_name': 'id_tipo_iva.codigo_iva', 'date_format': None},
		{'field_name': 'cuit', 'date_format': None},
		{'field_name': 'telefono_cliente', 'date_format': None},
	]


# ClienteListView - Inicio
class ClienteInformeListView(InformeListView):
	model = ConfigViews.model
	form_class = ConfigViews.form_class
	template_name = ConfigViews.template_list
	context_object_name = ConfigViews.context_object_name
	
	search_fields = DataViewList.search_fields
	ordering = DataViewList.ordering
	
	extra_context = {
		"master_title": f'Informes - {ConfigViews.model._meta.verbose_name_plural}',
		"home_view_name": ConfigViews.home_view_name,
		"list_view_name": ConfigViews.list_view_name,
		"table_headers": DataViewList.table_headers,
		"table_data": DataViewList.table_data,
		# "buscador_template": "informes/buscador_cliente.html",  # Se agrega el nombre de la plantilla buscador
		"buscador_template": f"{ConfigViews.app_label}/buscador_{ConfigViews.model_string}.html",  # Se agrega el nombre de la plantilla buscador,
		"js_file": ConfigViews.js_file
	}
	
	def get_queryset(self):
		queryset = super().get_queryset()
		form = self.form_class(self.request.GET)
		
		if form.is_valid():
			
			orden = form.cleaned_data.get('orden', 'nombre_cliente')
			desde = form.cleaned_data.get('desde', '').lower()
			hasta = form.cleaned_data.get('hasta', '').lower()
			vendedor = form.cleaned_data.get('vendedor')
			provincia = form.cleaned_data.get('provincia')
			localidad = form.cleaned_data.get('localidad')
			
			if orden not in ['nombre_cliente', 'id_cliente']:
				orden = 'nombre_cliente'
			
			queryset = queryset.order_by(orden)
			
			if orden == 'nombre_cliente':
				#-- Anotar un campo en minúsculas para la comparación insensible a mayúsculas/minúsculas.
				queryset = queryset.annotate(nombre_lower=Lower('nombre_cliente'))
				
				if desde and hasta:
					#-- Filtrar clientes cuyos nombres comienzan con letras en el rango desde-hasta.
					queryset = queryset.filter(
						Q(nombre_lower__gte=desde) &  # Nombres mayor o igual a "desde"
						Q(nombre_lower__lt=chr(ord(hasta[0]) + 1))  # Menor que la siguiente letra de "hasta"
					)
				elif desde:
					#-- Filtrar solo clientes mayores o iguales a "desde".
					queryset = queryset.filter(nombre_lower__gte=desde)
				elif hasta:
					#-- Filtrar solo clientes menores que la siguiente letra de "hasta".
					queryset = queryset.filter(nombre_lower__lt=chr(ord(hasta[0]) + 1))
				
				
			elif orden == 'id_cliente':
				if desde and hasta:
					queryset = queryset.filter(id_cliente__range=(desde, hasta))
				elif desde:
					queryset = queryset.filter(id_cliente__gte=desde)
				elif hasta:
					queryset = queryset.filter(id_cliente__lte=hasta)
			
			
			if vendedor:
				queryset = queryset.filter(id_vendedor=vendedor)
			
			if provincia:
				queryset = queryset.filter(id_provincia=provincia)
			
			if localidad:
				queryset = queryset.filter(id_localidad=localidad)
		else:
			#-- Agregar clases css a los campos con errores.
			print("El form no es válido (desde la vista)")
			print(f"{form.errors = }")
			form.add_error_classes()
						
		return queryset
	
	def get_context_data(self, **kwargs):
		context = super().get_context_data(**kwargs)
		form = BuscadorClienteForm(self.request.GET or None)
		
		context["form"] = form
		
		#-- Si el formulario tiene errores, pasa los errores al contexto.
		if form.errors:
			context["data_has_errors"] = True
		
		return context


class ClienteInformesView(View):
	"""Vista para gestionar informes de clientes, exportaciones y envíos por correo."""
	
	def get(self, request, *args, **kwargs):
		"""Gestión de solicitudes GET."""
		
		#-- "email" o "download".
		action = request.GET.get("action", "download")
		
		#-- Formatos seleccionados por el usuario.
		formatos = request.GET.getlist("formatos[]")
		
		#-- Email si aplica envío.
		email = request.GET.get("email", "")
		
		#-- Obtener el queryset filtrado.
		cliente_list_view = ClienteInformeListView()
		cliente_list_view.request = request
		queryset = cliente_list_view.get_queryset()
		
		#-- Generar y retornar el archivo ZIP.
		if action == "email":
			#-- Manejar el envío por correo electrónico.
			return self.enviar_por_email(queryset, formatos, email)
		else:
			#-- Manejar la generación y descarga del archivo ZIP.
			return self.generar_archivos_zip(queryset, formatos)
	
	def generar_archivos_zip(self, queryset, formatos):
		"""Generar un archivo ZIP con los formatos seleccionados."""
		
		buffer = BytesIO()
		with ZipFile(buffer, "w") as zip_file:
			helper = ExportHelper(queryset)
			
			#-- Generar los formatos seleccionados.
			if "pdf" in formatos:
				pdf_content = helper.export_to_pdf()
				zip_file.writestr("informe_clientes.pdf", pdf_content)
			
			if "csv" in formatos:
				csv_content = helper.export_to_csv()
				zip_file.writestr("informe_clientes.csv", csv_content)
			
			if "word" in formatos:
				word_content = helper.export_to_word()
				zip_file.writestr("informe_clientes.docx", word_content)
			
			if "excel" in formatos:
				excel_content = helper.export_to_excel()
				zip_file.writestr("informe_clientes.xlsx", excel_content)
		
		#-- Preparar respuesta para descargar el archivo ZIP.
		buffer.seek(0)
		response = HttpResponse(buffer, content_type="application/zip")
		response["Content-Disposition"] = 'attachment; filename="informe_clientes.zip"'
		
		return response
	
	def enviar_por_email(self, queryset, formatos, email):
		"""Enviar los informes seleccionados por correo electrónico."""
		helper = ExportHelper(queryset)
		attachments = []
		
		#-- Generar los formatos seleccionados y añadirlos como adjuntos.
		if "pdf" in formatos:
			attachments.append(("informe_clientes.pdf", helper.generar_pdf(), 
					   "application/pdf"))
		
		if "csv" in formatos:
			attachments.append(("informe_clientes.csv", helper.generar_csv(), 
					   "text/csv"))
		
		if "word" in formatos:
			attachments.append(("informe_clientes.docx", helper.generar_word(), 
					   "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))
		
		if "excel" in formatos:
			attachments.append(("informe_clientes.xlsx", helper.generar_excel(), 
					   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"))
		
		#-- Crear y enviar el correo.
		subject = "Informe de Clientes"
		body = "Adjunto encontrarás el informe solicitado."
		email_message = EmailMessage(subject, body, to=[email])
		for filename, content, mime_type in attachments:
			email_message.attach(filename, content, mime_type)
		
		email_message.send()

		#-- Responder con un mensaje de éxito.
		return JsonResponse({"success": True, "message": "Informe enviado correctamente al correo."})


class ClienteInformePDFView(View):
	
	def get(self, request, *args, **kwargs):
		#-- Obtener el queryset (el listado de clientes) ya filtrado.
		cliente_list_view = ClienteInformeListView()
		cliente_list_view.request = request
		queryset = cliente_list_view.get_queryset()
		
		#-- Generar el pdf.
		helper = ExportHelper(queryset)
		buffer = helper.export_to_pdf()
		
		#-- Preparar la respuesta HTTP.
		response = HttpResponse(buffer, content_type='application/pdf')
		response['Content-Disposition'] = 'inline; filename="informe_clientes.pdf"'
		
		return response


class ExportHelper:
	
	def __init__(self, queryset):
		self.queryset = queryset
	
	def _safe_str(self, value):
		return str(value) if value is not None else ""
	
	def export_to_pdf(self):
		buffer = BytesIO()
		doc = SimpleDocTemplate(buffer, pagesize=letter)
		elements = []
		
		# Encabezado
		elements.append(Paragraph("Reporte de Clientes", getSampleStyleSheet()['Title']))
		elements.append(Spacer(1, 12))
		
		# Tabla
		table_data = [[
			"Código", "Nombre", "Domicilio", "Localidad", 
			"C.P.", "IVA", "CUIT", "Teléfono"
		]]
		for obj in self.queryset:
			table_data.append([
				obj.id_cliente, obj.nombre_cliente, obj.domicilio_cliente,
				obj.id_localidad, obj.id_localidad.codigo_postal,
				obj.id_tipo_iva.codigo_iva, obj.cuit, obj.telefono_cliente
			])
		table = Table(table_data)
		table.setStyle(TableStyle([
			('BACKGROUND', (0, 0), (-1, 0), colors.grey),
			('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
			('ALIGN', (0, 0), (-1, -1), 'CENTER'),
			('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
			('BOTTOMPADDING', (0, 0), (-1, 0), 12),
			('BACKGROUND', (0, 1), (-1, -1), colors.beige),
			('GRID', (0, 0), (-1, -1), 1, colors.black),
		]))
		elements.append(table)
		
		doc.build(elements)
		buffer.seek(0)
		
		return buffer.getvalue()
	
	def export_to_csv(self):
		#-- Crear buffer binario.
		buffer = BytesIO()
		
		#-- Crear envoltura de texto.
		text_buffer = TextIOWrapper(buffer, encoding="utf-8", newline="")
		
		#-- Escribir el BOM manualmente para indicar UTF-8.
		text_buffer.write('\ufeff')  # Agregar el BOM al archivo como texto.
		
		writer = csv.writer(text_buffer)
		writer.writerow([
			"Código", "Nombre", "Domicilio", "Localidad",
			"C.P.", "IVA", "CUIT", "Teléfono"
		])
		for obj in self.queryset:
			writer.writerow([
				obj.id_cliente, obj.nombre_cliente, obj.domicilio_cliente,
				obj.id_localidad, obj.id_localidad.codigo_postal,
				obj.id_tipo_iva.codigo_iva, obj.cuit, obj.telefono_cliente
			])
		#-- Vaciar contenido al buffer binario.
		text_buffer.flush()
		buffer.seek(0)
		
		#-- Obtener el contenido en bytes.
		csv_bytes = buffer.getvalue()		
		
		#-- Cerrar buffers.
		text_buffer.close()
		buffer.close()
		
		return csv_bytes
	
	def export_to_word(self):
		
		doc = Document()
		doc.add_heading("Reporte de Clientes", level=1)
		
		table = doc.add_table(rows=1, cols=8)
		hdr_cells = table.rows[0].cells
		headers = [
			"Código", "Nombre", "Domicilio", "Localidad",
			"C.P.", "IVA", "CUIT", "Teléfono"
		]
		for i, header in enumerate(headers):
			hdr_cells[i].text = header
		
		for obj in self.queryset:
			row_cells = table.add_row().cells
			row_cells[0].text = self._safe_str(obj.id_cliente)
			row_cells[1].text = self._safe_str(obj.nombre_cliente)
			row_cells[2].text = self._safe_str(obj.domicilio_cliente)
			row_cells[3].text = self._safe_str(obj.id_localidad)
			row_cells[4].text = self._safe_str(obj.id_localidad.codigo_postal)
			row_cells[5].text = self._safe_str(obj.id_tipo_iva.codigo_iva)
			row_cells[6].text = self._safe_str(obj.cuit)
			row_cells[7].text = self._safe_str(obj.telefono_cliente)
		
		buffer = BytesIO()
		doc.save(buffer)
		buffer.seek(0)
		
		return buffer.getvalue()
	
	def export_to_excel(self):
		wb = Workbook()
		ws = wb.active
		ws.title = "Clientes"
		
		#-- Encabezados.
		headers = [
			"Código", "Nombre", "Domicilio", "Localidad",
			"C.P.", "IVA", "CUIT", "Teléfono"
		]
		ws.append(headers)
		
		#-- Datos.
		for obj in self.queryset:
			ws.append([
				obj.id_cliente, 
				self._safe_str(obj.nombre_cliente), 
				self._safe_str(obj.domicilio_cliente),
				self._safe_str(obj.id_localidad), 
				self._safe_str(obj.id_localidad.codigo_postal),
				self._safe_str(obj.id_tipo_iva.codigo_iva), 
				self._safe_str(obj.cuit), 
				self._safe_str(obj.telefono_cliente)
			])
		
		buffer = BytesIO()
		wb.save(buffer)
		buffer.seek(0)
		
		return buffer.getvalue()
