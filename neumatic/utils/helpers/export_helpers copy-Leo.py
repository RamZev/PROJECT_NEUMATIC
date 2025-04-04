# neumatic\utils\helpers\export_helpers.py

from io import BytesIO, TextIOWrapper
from reportlab.platypus import BaseDocTemplate, SimpleDocTemplate, Frame, PageTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from docx import Document
from openpyxl import Workbook
import csv
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from datetime import datetime
from decimal import Decimal
from apps.maestros.templatetags.custom_tags import formato_es_ar


class ExportHelper:
	
	def __init__(self, queryset, table_headers, report_title, total_columns=None):
		self.queryset = queryset
		self.table_headers = table_headers
		self.report_title = report_title
		self.total_columns = total_columns if total_columns else {}  # Diccionario con texto y columnas a totalizar
	
	def _safe_str(self, value):
		return str(value) if value is not None else ""
	
	def _get_headers_and_fields(self):
		"""Obtener nombres de encabezados y los campos correspondientes del diccionario table_headers."""
		
		headers = [header[1] for header in self.table_headers.values()]
		fields = list(self.table_headers.keys())
		return headers, fields
	
	def _resolve_field(self, obj, field_name, frmto=None):
		"""
		Resuelve el valor de un campo (incluso anidado).
		Si frmto es 'pdf', se formatea el valor (por ejemplo, números se convierten a cadena formateada);
		en caso contrario, se devuelve el valor en su tipo original (para Excel/CSV) para que los datos numéricos
		sean numéricos (enteros y Decimal con 2 decimales).
		"""
		try:
			#-- Recorrer los niveles anidados para obtener el valor.
			fields = field_name.split('.')
			value = obj
			for f in fields:
				#-- Si el atributo existe, se obtiene.
				if hasattr(value, f):
					value = getattr(value, f)
				#-- Sino, se intenta obtener del diccionario interno (__dict__).
				elif isinstance(value, dict) and f in value:
					value = value[f]
				else:
					value = None
					break
					
			if value is None:
				return ""
			
			#-- Si es booleano.
			if isinstance(value, bool):
				if frmto == 'pdf':
					#-- Para PDF, se formatea según si el campo contiene "estatus".
					if "estatus" in fields[-1]:
						return "Activo" if value else "Inactivo"
					else:
						return "Sí" if value else "No"
				else:
					return value
			
			#-- Si es numérico.
			if isinstance(value, (float, Decimal)):
				if frmto == 'pdf':
					#-- Para PDF, formatea el número a cadena.
					return self._safe_str(formato_es_ar(value))
				else:
					#-- Para Excel/CSV, devuelve el número:
					#-- - Si es float, redondea a 2 decimales.
					#-- - Si es Decimal, lo cuantiza a 2 decimales.
					if isinstance(value, float):
						return round(value, 2)
					else:
						return value.quantize(Decimal('0.01'))
			
			#-- Otros tipos.
			if frmto == 'pdf':
				return self._safe_str(value)
			else:
				return value
		except Exception as e:
			return ""
	
	def _calculate_totals(self, fields):
		"""Calcula los totales para las columnas especificadas en total_columns y genera la fila de totales."""
		
		totals = [""] * len(fields)  #-- Inicializa la fila de totales vacía.
		
		for total_text, columns in self.total_columns.items():
			#-- Encuentra la primera columna para colocar el texto.
			first_col = None
			for col in columns:
				if col in fields:
					col_index = fields.index(col)
					total = sum(
						getattr(obj, col.split('.')[0], 0) or 0
						for obj in self.queryset
					)
					#-- Formatear el total.
					totals[col_index] = formato_es_ar(total)
					
					if first_col is None:  # Coloca el texto solo una vez
						first_col = col_index
			
			#-- Coloca el texto en la primera columna encontrada.
			if first_col is not None:
				totals[first_col - 1 if first_col > 0 else 0] = total_text
		
		return totals
	
	def export_to_pdf(self):
		"""Exporta los datos del queryset a un archivo PDF."""
		
		buffer = BytesIO()
		doc = SimpleDocTemplate(buffer, pagesize=A4)
		elements = []
		
		#-- Encabezado del reporte.
		elements.append(Paragraph(self.report_title, getSampleStyleSheet()['Title']))
		elements.append(Spacer(1, 12))
		
		#-- Obtener encabezados y campos dinámicos.
		headers, fields = self._get_headers_and_fields()
		
		#-- Crear la tabla de datos.
		table_data = [headers]  #-- Primera fila con los encabezados.
		
		#-- Agregar los datos desde el queryset.
		for obj in self.queryset:
			table_data.append([self._resolve_field(obj, field, frmto='pdf') for field in fields])
		
		#-- Calcular y agregar la fila de totales.
		if self.total_columns:
			totals_row = self._calculate_totals(fields)
			table_data.append(totals_row)
		
		table = Table(table_data)
		table_style = TableStyle([
			('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
			('FONTSIZE', (0, 0), (-1, -1), 8),
			# ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
			('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
			# Línea horizontal por encima de los títulos
			('LINEABOVE', (0, 0), (-1, 0), 1, colors.black),
			# Línea horizontal por debajo de los títulos
			('LINEBELOW', (0, 0), (-1, 0), 1, colors.black),
			('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
			
			# ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
			# ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
			# ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
			# ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
			# ('GRID', (0, 0), (-1, -1), 1, colors.black),
			
		])
			
		#-- Identificar las columnas que contienen números para alinear a la derecha.
		for col_idx, field in enumerate(fields):
			if any(isinstance(getattr(obj, field.split('.')[0], None), (int, float, Decimal)) for obj in self.queryset):
				table_style.add('ALIGN', (col_idx, 1), (col_idx, -1), 'RIGHT')
		
		#-- Alinear la fila de totales y aplicar negritas.
		if self.total_columns:
			totals_row_index = len(table_data) - 1  # Índice de la fila de totales (última fila)
			
			#-- Negritas en toda la fila de totales.
			table_style.add('FONTNAME', (0, totals_row_index), (-1, totals_row_index), 'Helvetica-Bold')
			
			#-- Agregar línea horizontal encima de la fila de totales.
			table_style.add('LINEABOVE', (0, -1), (-1, -1), 1, colors.black)
			
			#-- Alinear columnas numéricas a la derecha en la fila de totales.
			for total_text, columns in self.total_columns.items():
				for col in columns:
					if col in fields:
						col_index = fields.index(col)
						table_style.add('ALIGN', (col_index, -1), (col_index, -1), 'RIGHT')
			
		table.setStyle(table_style)
		
		#-- Repite la primera fila (headers) en cada página.
		table.repeatRows = 1
		
		elements.append(table)
		
		doc.build(elements, canvasmaker=CustomCanvas)
		buffer.seek(0)
		
		return buffer.getvalue()
	
	def export_to_csv(self):
		"""Exporta los datos del queryset a un archivo CSV."""
		
		#-- Crear buffer binario.
		buffer = BytesIO()
		
		#-- Crear envoltura de texto.
		text_buffer = TextIOWrapper(buffer, encoding="utf-8", newline="")
		
		#-- Escribir el BOM manualmente para indicar UTF-8.
		text_buffer.write('\ufeff')  #-- Agregar el BOM al archivo como texto.
		
		writer = csv.writer(text_buffer)
		
		headers, fields = self._get_headers_and_fields()
		writer.writerow(headers)  #-- Escribir encabezados.
		
		for obj in self.queryset:
			row = [self._resolve_field(obj, field) for field in fields]
			writer.writerow(row)
		
		#-- Vaciar contenido al buffer binario.
		text_buffer.flush()
		buffer.seek(0)
		
		#-- Obtener el contenido en bytes.
		csv_bytes = buffer.getvalue()		
		
		#-- Cerrar buffers.
		text_buffer.close()
		buffer.close()
		
		return csv_bytes
	
	def export_to_word(self):
		"""Exporta el queryset a un documento de Word."""
		
		#-- Crear el documento.
		doc = Document()
		
		#-- Agregar un título.
		doc.add_heading(self.report_title, level=1)
		
		#-- Obtener encabezados y campos.
		headers, fields = self._get_headers_and_fields()
		
		#-- Crear tabla con encabezados dinámicos.
		table = doc.add_table(rows=1, cols=len(headers))
		hdr_cells = table.rows[0].cells
		
		#-- Agregar encabezados a la tabla.
		for i, header in enumerate(headers):
			hdr_cells[i].text = header
		
		#-- Agregar las filas de datos.
		for obj in self.queryset:
			row_cells = table.add_row().cells
			for i, field in enumerate(fields):
				value = self._resolve_field(obj, field)
				row_cells[i].text = value
		
		buffer = BytesIO()
		doc.save(buffer)
		buffer.seek(0)
		
		return buffer.getvalue()
	
	def export_to_excel(self):
		"""Exporta los datos del queryset a un archivo Excel."""
		
		wb = Workbook()
		ws = wb.active
		ws.title = self.report_title
		
		#-- Obtener encabezados y campos dinámicos.
		headers, fields = self._get_headers_and_fields()
		
		#-- Encabezados.
		ws.append(headers)
		
		#-- Datos.
		for obj in self.queryset:
			row = [self._resolve_field(obj, field) for field in fields]
			ws.append(row)
		
		buffer = BytesIO()
		wb.save(buffer)
		buffer.seek(0)
		
		return buffer.getvalue()


class CustomCanvas(canvas.Canvas):
	def __init__(self, *args, **kwargs):
		super().__init__(*args, **kwargs)
		
		#-- Contador para el número de la página actual.
		self.page_number = 0
		
		#-- Contador para el total de páginas.
		self.total_pages = 0
		
		#-- Fecha del reporte y su formato.
		# self.report_date = datetime.now().strftime("Fecha: %d/%m/%Y")
		self.report_date = datetime.now().strftime("%d/%m/%Y %I:%M %p")
	
	def draw_footer(self):
		# Dibujar el pie de página con tres secciones

		# Margen inferior
		y_position = 20
		
		# Extremo izquierdo: texto fijo
		self.setFont("Helvetica", 8)
		self.setFillColor(colors.black)
		
		self.drawString(30, y_position, "M.A.A.Soft")
		
		#-- Incrementar el número total de páginas.
		self.page_number += 1
		
		# Centro: numeración de páginas
		# page_text = f"Página {self.page_number}/{self.total_pages}"
		page_text = f"- {self.page_number} -"
		self.drawCentredString(300, y_position, page_text)
		
		# Extremo derecho: fecha del reporte
		self.drawRightString(570, y_position, self.report_date)
		
		# Línea horizontal por encima del pie de página
		self.line(30, y_position + 10, 570, y_position + 10)
	
	def showPage(self):
		"""Llama a la función que dibuja elementos comunes en cada página."""
		
		#-- Incrementar el número total de páginas cuando se muestra una nueva página.
		self.total_pages += 1
		
		#-- Dibujar el pie de página antes de pasar a la siguiente página.
		self.draw_footer()
		
		#-- Llamar al método showPage original para cambiar de página.
		super().showPage()


#---------------------------------------------------------------------------------
'''
class PDFGenerator:
	def __init__(self, context, pagesize=landscape(A4), margins=(10, 10, 0, 40)):
		self.buffer = BytesIO()
		self.context = context
		self.pagesize = pagesize
		self.left_margin, self.right_margin, self.top_margin, self.bottom_margin = margins
		
		#-- Creación y Configuración del documento.
		self.doc = BaseDocTemplate(
			self.buffer,
			pagesize=self.pagesize,
			leftMargin=self.left_margin,
			rightMargin=self.right_margin,
			topMargin=self.top_margin,
			bottomMargin=self.bottom_margin
		)
		
		#-- Crear y Configurar frame.
		frame = Frame(
			self.doc.leftMargin,
			self.doc.bottomMargin,
			self.doc.width,
			self.doc.height - 80,  # Ajustar según necesidades
			id="body"
		)
		
		#-- Crear y Configurar template.
		template = PageTemplate(
			id="reportTemplate", 
			frames=[frame], 
			onPage=self._header_footer
		)
		self.doc.addPageTemplates([template])
		self.doc.contexto_reporte = self.context
		
		#-- Estilos.
		self.styles = getSampleStyleSheet()
		self._setup_custom_styles()
	
	def _setup_custom_styles(self):
		#-- Estilos personalizados.
		
		#-- Estilo de párrafo para los datos en la tabla del reporte.
		self.styles.add(ParagraphStyle(
			name='CellStyle',
			parent=self.styles["BodyText"],
			fontSize=6,
			leading=5.0,
			spaceBefore=0,
			spaceAfter=0,
			leftIndent=0,
			rightIndent=0,
			firstLineIndent=0,
		))
		
		#-- Estilo de párrafo para la sección Header-bottom-left.
		self.styles.add(ParagraphStyle(
			name='HeaderBottomLeft',
			fontSize=9,
			leading=12,
			alignment=TA_LEFT
		))
		
		#-- Estilo de párrafo para la sección Header-bottom-right.
		self.styles.add(ParagraphStyle(
			name='HeaderBottomRight',
			fontSize=9,
			leading=12,
			alignment=TA_RIGHT
		))
	
	def _header_footer(self, canvas_obj, doc):
		"""Método para dibujar header y footer en cada página"""
		canvas_obj.saveState()
		width, height = doc.pagesize
		
		#-----------------------------------------------------------------------------------
		#  HEADER
		#-----------------------------------------------------------------------------------
		
		header_top_height = 50
		
		# --- Header-top -------------------------------------------------------------------
		
		#-- Sección Superior Izquierda: logotipo.
		logo_path = doc.contexto_reporte.get("logo_url", "")
		logo_area_width = (width - doc.leftMargin - doc.rightMargin) * 0.25
		logo_height = 30
		logo_x = doc.leftMargin - 30
		logo_y = height - header_top_height + (header_top_height - logo_height) / 2
		
		try:
			canvas_obj.drawImage(
				logo_path, logo_x, logo_y, 
				width=logo_area_width, height=logo_height, 
				preserveAspectRatio=True, mask='auto'
			)
		except Exception:
			canvas_obj.setFont("Helvetica", 10)
			canvas_obj.drawString(logo_x, logo_y, "[Logo]")
		
		#-- Sección Superior Perecha: título.
		titulo = doc.contexto_reporte.get("titulo", "Reporte")
		canvas_obj.setFont("Helvetica-BoldOblique", 12)
		canvas_obj.drawRightString(
			width - doc.rightMargin, 
			(height - header_top_height/2)-10, 
			titulo
		)
		
		# --- Header-bottom ----------------------------------------------------------------
		
		#-- Línea superior de header-bottom.
		line_y_start = height - header_top_height  
		canvas_obj.setLineWidth(1)
		canvas_obj.setStrokeColor(colors.black)
		canvas_obj.line(doc.leftMargin, line_y_start, width - doc.rightMargin, line_y_start)
		
		#-- Sección inferior Izquierda.
		header_bottom_text_left = doc.contexto_reporte.get("header_bottom_left", "")
		
		#-- Sección inferior Derecha.
		parametros = doc.contexto_reporte.get("parametros", {})
		lines = [f"<b>{key}:</b> {value}" for key, value in parametros.items()]
		header_bottom_text_right = "<br/>".join(lines)
		
		#-- Definir estilos y posición para cada sección.
		p_left = Paragraph(header_bottom_text_left, self.styles['HeaderBottomLeft'])
		p_right = Paragraph(header_bottom_text_right, self.styles['HeaderBottomRight'])
		
		available_width = (width - doc.leftMargin - doc.rightMargin) / 2.0
		left_w, left_h = p_left.wrap(available_width, 100)
		right_w, right_h = p_right.wrap(available_width, 100)
		header_bottom_height = max(left_h, right_h)
		header_bottom_y = height - header_top_height - header_bottom_height
		
		p_left.drawOn(canvas_obj, doc.leftMargin, header_bottom_y)
		p_right.drawOn(canvas_obj, doc.leftMargin + available_width, header_bottom_y)
		
		#-- Línea inferior de header-bottom.
		canvas_obj.line(doc.leftMargin, header_bottom_y, width - doc.rightMargin, header_bottom_y)
		
		#-----------------------------------------------------------------------------------
		#  Footer
		#-----------------------------------------------------------------------------------
		footer_y = 15
		
		#-- Línea decorativa.
		line_y = footer_y + 12
		canvas_obj.setLineWidth(1)
		canvas_obj.setStrokeColor(colors.black)
		canvas_obj.line(doc.leftMargin, line_y, width - doc.rightMargin, line_y)
		
		#-- Texto a la izquierda del footer.
		canvas_obj.setFont("Helvetica-Oblique", 9)
		canvas_obj.drawString(doc.leftMargin, footer_y, "M.A.A.Soft")
		
		#-- Número de página formateado al centro.
		numero_pagina_text = f"Página {canvas_obj._pageNumber}"
		canvas_obj.setFont("Helvetica", 9)
		canvas_obj.drawCentredString(width/2.0, footer_y, numero_pagina_text)
		
		#-- Fecha y hora del reporte a la derecha.
		fecha_reporte = datetime.now().strftime("%d/%m/%Y %H:%M")
		canvas_obj.setFont("Helvetica", 9)
		canvas_obj.drawRightString(width - doc.rightMargin, footer_y, fecha_reporte)
		
		canvas_obj.restoreState()
	
	def _create_table(self, data, col_widths, style_config):
		"""Crea una tabla con los datos y estilos proporcionados"""
		table = Table(data, colWidths=col_widths, repeatRows=1)
		
		#-- Estilo base.
		table_style = TableStyle([
			('BACKGROUND', (0,0), (-1,0), colors.gray),
			('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
			('VALIGN', (0,0), (-1,-1), 'TOP'),
			('FONTSIZE', (0,0), (-1,-1), 6),
			('TOPPADDING', (0,1), (-1,1), 2),
			('TOPPADDING', (0,2), (-1,-1), 0),
			('BOTTOMPADDING', (0,1), (-1,-1), 0),
		])
		
		#-- Añadir configuraciones de estilo adicionales.
		for config in style_config:
			table_style.add(*config)
		
		table.setStyle(table_style)
		return table
	
	def generate(self, table_data, col_widths, table_style_config):
		"""Genera el PDF con los datos proporcionados"""
		content = []
		
		#-- Crear tabla.
		table = self._create_table(table_data, col_widths, table_style_config)
		content.append(table)
		
		#-- Construir documento.
		self.doc.build(content)
		
		pdf = self.buffer.getvalue()
		self.buffer.close()
		return pdf
'''

class PDFGenerator:
	def __init__(self, context, pagesize=landscape(A4), margins=(10, 10, 0, 40)):
		self.buffer = BytesIO()
		self.context = context
		self.pagesize = pagesize
		self.left_margin, self.right_margin, self.top_margin, self.bottom_margin = margins
		
		#-- Creación y Configuración del documento.
		self.doc = BaseDocTemplate(
			self.buffer,
			pagesize=self.pagesize,
			leftMargin=self.left_margin,
			rightMargin=self.right_margin,
			topMargin=self.top_margin,
			bottomMargin=self.bottom_margin
		)
		
		#-- Crear y Configurar frame.
		frame = Frame(
			self.doc.leftMargin,
			self.doc.bottomMargin,
			self.doc.width,
			self.doc.height - 80,  # Ajustar según necesidades
			id="body"
		)
		
		#-- Crear y Configurar template.
		template = PageTemplate(
			id="reportTemplate", 
			frames=[frame], 
			onPage=self._header_footer
		)
		self.doc.addPageTemplates([template])
		self.doc.contexto_reporte = self.context
		
		#-- Estilos.
		self.styles = getSampleStyleSheet()
		self._setup_custom_styles()
	
	def _setup_custom_styles(self):
		#-- Estilos personalizados.
		
		#-- Estilo de párrafo para los datos en la tabla del reporte.
		self.styles.add(ParagraphStyle(
			name='CellStyle',
			parent=self.styles["BodyText"],
			fontSize=6,
			leading=5.0,
			spaceBefore=0,
			spaceAfter=0,
			leftIndent=0,
			rightIndent=0,
			firstLineIndent=0,
		))
		
		#-- Estilo de párrafo para la sección Header-bottom-left.
		self.styles.add(ParagraphStyle(
			name='HeaderBottomLeft',
			fontSize=9,
			leading=12,
			alignment=TA_LEFT
		))
		
		#-- Estilo de párrafo para la sección Header-bottom-right.
		self.styles.add(ParagraphStyle(
			name='HeaderBottomRight',
			fontSize=9,
			leading=12,
			alignment=TA_RIGHT
		))
	
	# ----------------------------------------------------------------------------------
	def _render_header_bottom(self, canvas_obj, doc, width, height):
		"""Nuevo método modular para el header inferior"""
		#-- Línea superior.
		line_y_start = height - 50  
		canvas_obj.setLineWidth(1)
		canvas_obj.line(doc.leftMargin, line_y_start, width - doc.rightMargin, line_y_start)
		
		#-- Contenido personalizable.
		left_content = self._get_header_bottom_left(doc.contexto_reporte)
		right_content = self._get_header_bottom_right(doc.contexto_reporte)
		
		#-- Renderizado.
		self._draw_header_bottom_content(
			canvas_obj, doc, 
			left_content, right_content,
			line_y_start
		)
	
	def _get_header_bottom_left(self, context):
		"""SOBREESCRIBIR ESTE MÉTODO PARA CONTENIDO IZQUIERDO PERSONALIZADO"""
		return context.get("header_bottom_left", "")
	
	def _get_header_bottom_right(self, context):
		"""SOBREESCRIBIR ESTE MÉTODO PARA CONTENIDO DERECHO PERSONALIZADO"""
		params = context.get("parametros", {})
		return "<br/>".join([f"<b>{k}:</b> {v}" for k, v in params.items()])
	
	def _draw_header_bottom_content(self, canvas_obj, doc, left_content, right_content, start_y):
		"""Lógica compartida para renderizar ambos lados"""
		p_left = Paragraph(left_content, self.styles['HeaderBottomLeft'])
		p_right = Paragraph(right_content, self.styles['HeaderBottomRight'])
		
		available_width = (doc.width - doc.leftMargin - doc.rightMargin) / 2.0
		left_w, left_h = p_left.wrap(available_width, 100)
		right_w, right_h = p_right.wrap(available_width, 100)
		content_height = max(left_h, right_h)
		content_y = start_y - content_height
		
		p_left.drawOn(canvas_obj, doc.leftMargin, content_y)
		p_right.drawOn(canvas_obj, doc.leftMargin + available_width, content_y)
		
		#-- Línea inferior.
		canvas_obj.line(doc.leftMargin, content_y, doc.width - doc.rightMargin, content_y)
	# ----------------------------------------------------------------------------------
	
	def _header_footer(self, canvas_obj, doc):
		"""Método para dibujar header y footer en cada página"""
		canvas_obj.saveState()
		width, height = doc.pagesize
		
		#-- Header Top (logo + título).
		self._render_header_top(canvas_obj, doc, width, height)
		
		#-- Header Bottom.
		self._render_header_bottom(canvas_obj, doc, width, height)
		
		#-- Footer.
		self._render_footer(canvas_obj, doc, width)
		
		canvas_obj.restoreState()
		
	def _render_header_top(self, canvas_obj, doc, width, height):
		# --- Header-top -------------------------------------------------------------------
		header_top_height = 50
		
		#-- Sección Superior Izquierda: logotipo.
		logo_path = doc.contexto_reporte.get("logo_url", "")
		logo_area_width = (width - doc.leftMargin - doc.rightMargin) * 0.25
		logo_height = 30
		logo_x = doc.leftMargin - 30
		logo_y = height - header_top_height + (header_top_height - logo_height) / 2
		
		try:
			canvas_obj.drawImage(
				logo_path, logo_x, logo_y, 
				width=logo_area_width, height=logo_height, 
				preserveAspectRatio=True, mask='auto'
			)
		except Exception:
			canvas_obj.setFont("Helvetica", 10)
			canvas_obj.drawString(logo_x, logo_y, "[Logo]")
		
		#-- Sección Superior Perecha: título.
		titulo = doc.contexto_reporte.get("titulo", "Reporte")
		canvas_obj.setFont("Helvetica-BoldOblique", 12)
		canvas_obj.drawRightString(
			width - doc.rightMargin, 
			(height - header_top_height/2)-10, 
			titulo
		)
		
	def _render_footer(self, canvas_obj, doc, width):
		# --- Footer -----------------------------------------------------------------------
		footer_y = 15
		
		#-- Línea decorativa.
		line_y = footer_y + 12
		canvas_obj.setLineWidth(1)
		canvas_obj.setStrokeColor(colors.black)
		canvas_obj.line(doc.leftMargin, line_y, width - doc.rightMargin, line_y)
		
		#-- Texto a la izquierda del footer.
		canvas_obj.setFont("Helvetica-Oblique", 9)
		canvas_obj.drawString(doc.leftMargin, footer_y, "M.A.A.Soft")
		
		#-- Número de página formateado al centro.
		numero_pagina_text = f"Página {canvas_obj._pageNumber}"
		canvas_obj.setFont("Helvetica", 9)
		canvas_obj.drawCentredString(width/2.0, footer_y, numero_pagina_text)
		
		#-- Fecha y hora del reporte a la derecha.
		fecha_reporte = datetime.now().strftime("%d/%m/%Y %H:%M")
		canvas_obj.setFont("Helvetica", 9)
		canvas_obj.drawRightString(width - doc.rightMargin, footer_y, fecha_reporte)
	
	def _create_table(self, data, col_widths, style_config):
		"""Crea una tabla con los datos y estilos proporcionados"""
		table = Table(data, colWidths=col_widths, repeatRows=1)
		
		#-- Estilo base.
		table_style = TableStyle([
			('BACKGROUND', (0,0), (-1,0), colors.gray),
			('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
			('VALIGN', (0,0), (-1,-1), 'TOP'),
			('FONTSIZE', (0,0), (-1,-1), 6),
			('TOPPADDING', (0,1), (-1,1), 2),
			('TOPPADDING', (0,2), (-1,-1), 0),
			('BOTTOMPADDING', (0,1), (-1,-1), 0),
		])
		
		#-- Añadir configuraciones de estilo adicionales.
		for config in style_config:
			table_style.add(*config)
		
		table.setStyle(table_style)
		return table
	
	def generate(self, table_data, col_widths, table_style_config):
		"""Genera el PDF con los datos proporcionados"""
		content = []
		
		#-- Crear tabla.
		table = self._create_table(table_data, col_widths, table_style_config)
		content.append(table)
		
		#-- Construir documento.
		self.doc.build(content)
		
		pdf = self.buffer.getvalue()
		self.buffer.close()
		return pdf