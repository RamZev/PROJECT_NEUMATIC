# neumatic\utils\helpers\export_helpers.py

from reportlab.lib.pagesizes import A4
from io import BytesIO, TextIOWrapper
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from docx import Document
from openpyxl import Workbook
import csv
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from datetime import datetime
from decimal import Decimal
from apps.maestros.templatetags.custom_tags import formato_es_ar


class ExportHelper:
	
	def __init__(self, queryset, table_headers, report_title):
		self.queryset = queryset
		self.table_headers = table_headers
		self.report_title = report_title
	
	def _safe_str(self, value):
		return str(value) if value is not None else ""
	
	def _get_headers_and_fields(self):
		"""Obtener nombres de encabezados y los campos correspondientes del diccionario table_headers."""
		
		headers = [header[1] for header in self.table_headers.values()]
		fields = list(self.table_headers.keys())
		return headers, fields
	
	def _resolve_field(self, obj, field_name):
		"""Resuelve el valor de un campo, incluso si es un campo anidado."""
		
		try:
			fields = field_name.split('.')  #-- Separar los niveles de anidación.
			value = obj
			
			for field in fields:
				value = getattr(value, field, None)  #-- Obtener el siguiente nivel.
				
				if isinstance(value, bool) and "estatus" in field:
					value = "Activo" if value else "Inactivo"
				elif isinstance(value, bool):
					value = "Sí" if value else "No"
				elif isinstance(value, (float, Decimal)):
					value = formato_es_ar(value)  # Aplica el formato de número				
					
			return self._safe_str(value)  #-- Convertir el valor a una cadena segura.
		except AttributeError:
			return ""  #-- Si no se puede resolver el campo, devolver una cadena vacía.
	
	def export_to_pdf(self):
		"""Exporta los datos del queryset a un archivo PDF."""
		
		buffer = BytesIO()
		doc = SimpleDocTemplate(buffer, pagesize=A4)
		elements = []
		
		#-- Encabezado del reporte.
		elements.append(Paragraph(self.report_title, getSampleStyleSheet()['Title']))
		elements.append(Spacer(1, 12))
		
		#-- Obtener encabezados y campos dinámicos.
		headers, fields = self._get_headers_and_fields()
		
		#-- Crear la tabla de datos.
		table_data = [headers]  #-- Primera fila con los encabezados.
		
		#-- Agregar los datos desde el queryset.
		for obj in self.queryset:
			table_data.append([self._resolve_field(obj, field) for field in fields])
		
		table = Table(table_data)
		table_style = TableStyle([
			('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
			('FONTSIZE', (0, 0), (-1, -1), 8),
			# ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
			('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
			# Línea horizontal por encima de los títulos
			('LINEABOVE', (0, 0), (-1, 0), 1, colors.black),
			# Línea horizontal por debajo de los títulos
			('LINEBELOW', (0, 0), (-1, 0), 1, colors.black),
			('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
			
			# ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
			# ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
			# ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
			# ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
			# ('GRID', (0, 0), (-1, -1), 1, colors.black),
			
		])
			
		#-- Identificar las columnas que contienen números para alinear a la derecha.
		for col_idx, field in enumerate(fields):
			if any(isinstance(getattr(obj, field.split('.')[0], None), (int, float, Decimal)) for obj in self.queryset):
				table_style.add('ALIGN', (col_idx, 1), (col_idx, -1), 'RIGHT')
		
		table.setStyle(table_style)
		
		#-- Repite la primera fila (headers) en cada página.
		table.repeatRows = 1
		
		elements.append(table)
		
		doc.build(elements, canvasmaker=CustomCanvas)
		buffer.seek(0)
		
		return buffer.getvalue()
	
	def export_to_csv(self):
		"""Exporta los datos del queryset a un archivo CSV."""
		
		#-- Crear buffer binario.
		buffer = BytesIO()
		
		#-- Crear envoltura de texto.
		text_buffer = TextIOWrapper(buffer, encoding="utf-8", newline="")
		
		#-- Escribir el BOM manualmente para indicar UTF-8.
		text_buffer.write('\ufeff')  #-- Agregar el BOM al archivo como texto.
		
		writer = csv.writer(text_buffer)
		
		headers, fields = self._get_headers_and_fields()
		writer.writerow(headers)  #-- Escribir encabezados.
		
		for obj in self.queryset:
			row = [self._resolve_field(obj, field) for field in fields]
			writer.writerow(row)
		
		#-- Vaciar contenido al buffer binario.
		text_buffer.flush()
		buffer.seek(0)
		
		#-- Obtener el contenido en bytes.
		csv_bytes = buffer.getvalue()		
		
		#-- Cerrar buffers.
		text_buffer.close()
		buffer.close()
		
		return csv_bytes
	
	def export_to_word(self):
		"""Exporta el queryset a un documento de Word."""
		
		#-- Crear el documento.
		doc = Document()
		
		#-- Agregar un título.
		doc.add_heading(self.report_title, level=1)
		
		#-- Obtener encabezados y campos.
		headers, fields = self._get_headers_and_fields()
		
		#-- Crear tabla con encabezados dinámicos.
		table = doc.add_table(rows=1, cols=len(headers))
		hdr_cells = table.rows[0].cells
		
		#-- Agregar encabezados a la tabla.
		for i, header in enumerate(headers):
			hdr_cells[i].text = header
		
		#-- Agregar las filas de datos.
		for obj in self.queryset:
			row_cells = table.add_row().cells
			for i, field in enumerate(fields):
				value = self._resolve_field(obj, field)
				row_cells[i].text = value
		
		buffer = BytesIO()
		doc.save(buffer)
		buffer.seek(0)
		
		return buffer.getvalue()
	
	def export_to_excel(self):
		"""Exporta los datos del queryset a un archivo Excel."""
		
		wb = Workbook()
		ws = wb.active
		ws.title = self.report_title
		
		#-- Obtener encabezados y campos dinámicos.
		headers, fields = self._get_headers_and_fields()
		
		#-- Encabezados.
		ws.append(headers)
		
		#-- Datos.
		for obj in self.queryset:
			row = [self._resolve_field(obj, field) for field in fields]
			ws.append(row)
		
		buffer = BytesIO()
		wb.save(buffer)
		buffer.seek(0)
		
		return buffer.getvalue()


class CustomCanvas(canvas.Canvas):
	def __init__(self, *args, **kwargs):
		super().__init__(*args, **kwargs)
		
		#-- Contador para el número de la página actual.
		self.page_number = 0
		
		#-- Contador para el total de páginas.
		self.total_pages = 0
		
		#-- Fecha del reporte y su formato.
		# self.report_date = datetime.now().strftime("Fecha: %d/%m/%Y")
		self.report_date = datetime.now().strftime("%d/%m/%Y %I:%M %p")
	
	def draw_footer(self):
		# Dibujar el pie de página con tres secciones

		# Margen inferior
		y_position = 20
		
		# Extremo izquierdo: texto fijo
		self.setFont("Helvetica", 8)
		self.setFillColor(colors.black)
		
		self.drawString(30, y_position, "M.A.A.Soft")
		
		#-- Incrementar el número total de páginas.
		self.page_number += 1
		
		# Centro: numeración de páginas
		# page_text = f"Página {self.page_number}/{self.total_pages}"
		page_text = f"- {self.page_number} -"
		self.drawCentredString(300, y_position, page_text)
		
		# Extremo derecho: fecha del reporte
		self.drawRightString(570, y_position, self.report_date)
		
		# Línea horizontal por encima del pie de página
		self.line(30, y_position + 10, 570, y_position + 10)
	
	def showPage(self):
		"""Llama a la función que dibuja elementos comunes en cada página."""
		
		#-- Incrementar el número total de páginas cuando se muestra una nueva página.
		self.total_pages += 1
		
		#-- Dibujar el pie de página antes de pasar a la siguiente página.
		self.draw_footer()
		
		#-- Llamar al método showPage original para cambiar de página.
		super().showPage()
	
	'''
	def draw_page_number(self):
		"""Dibuja el número de página en el pie de página y una línea horizontal sobre él."""
		
		#-- Incrementar el número total de páginas.
		self.page_number += 1
		
		#-- Texto para el número de página.
		text = f"Página {self.page_number}/{self.total_pages}"
		
		#-- Posiciones para la línea y el texto.
		line_y = 40  #-- Posición Y de la línea.
		text_y = line_y + 10  #-- Posición Y del número de página (justo debajo de la línea).
		
		#-- Dibuja la línea horizontal.
		self.setStrokeColor(colors.black)
		self.setLineWidth(1)
		self.line(50, line_y, 550, line_y)  #-- Línea horizontal (ajustar las coordenadas según sea necesario).

		self.setFont("Helvetica", 10)
		self.setFillColor(colors.black)
		self.drawString(500, 20, text)  #-- Puedes ajustar las coordenadas si es necesario.
	
	def draw_date(self):
		"""Dibuja la fecha del reporte en la esquina superior derecha."""
		#-- Obtener la fecha actual en el formato "Fecha: dd/mm/aaaa".
		current_date = datetime.now().strftime("Fecha: %d/%m/%Y")
		
		#-- Posición de la fecha (esquina superior derecha).
		self.setFont("Helvetica", 10)
		self.setFillColor(colors.black)
		self.drawRightString(550, 800, current_date)  #-- Ajustar las coordenadas según sea necesario.
	
	def showPage(self):
		"""Llama a la función que dibuja elementos comunes en cada página."""
		
		#-- Incrementar el número total de páginas cuando se muestra una nueva página.
		self.total_pages += 1
		
		#-- Llamar a la función para agregar el número de página.
		self.draw_page_number()
		
		#-- Dibuja la fecha en la esquina superior derecha.
		self.draw_date()
		
		#-- Llamar al método showPage original para cambiar de página.
		super().showPage()
	
	'''